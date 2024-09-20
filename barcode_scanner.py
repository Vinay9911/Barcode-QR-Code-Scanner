from pyzbar.pyzbar import decode
import treepoem
import cv2
import os
from docx import Document

# Explicitly set the path to Ghostscript binary
os.environ['PATH'] += os.pathsep + r'C:\Program Files\gs\gs10.03.1\bin'

# Create the Word document
document = Document()
document.add_heading('Decoded Barcodes And QR Codes', 0)

# Set to store decoded barcodes
decoded_barcodes = set()

def create_barcode(filename, barcode_type, data):
    image = treepoem.generate_barcode(barcode_type, data)
    image.convert('1').save(filename)

# Example usage
create_barcode('barcode1.png', 'code128', '1234567890')
create_barcode('barcode2.png', 'code128', '0987654321')

def decoder(frame):
    try:
        decoding_info = decode(frame)
    except Exception as e:
        print(f"Error decoding barcode: {e}")
        decoding_info = []
    return decoding_info

def detect_codes_in_real_time():
    cap = cv2.VideoCapture(1)  # Adjust index as needed
    if not cap.isOpened():
        print("Error: Could not open video stream.")
        return
    while True:
        ret, frame = cap.read()
        if not ret:
            break
        
        decoding_info = decoder(frame)
        for barcode in decoding_info:
            x, y, w, h = barcode.rect
            cv2.rectangle(frame, (x, y), (x + w, y + h), (255, 0, 0), 2)
            message = barcode.data.decode('utf-8')
            
            if message not in decoded_barcodes:
                decoded_barcodes.add(message)
                document.add_paragraph(message)
                document.save('decoded_barcodes.docx')
            
            cv2.putText(frame, message, (x, y - 10), cv2.FONT_HERSHEY_SIMPLEX, 0.5, (255, 0, 0), 2)
            print(message)
        
        cv2.imshow('Frame', frame)
        key = cv2.waitKey(1)
        if key == 27:  # Press 'ESC' to exit
            break
    cap.release()
    cv2.destroyAllWindows()

detect_codes_in_real_time()
